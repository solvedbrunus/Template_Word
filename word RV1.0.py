"""
Word Template Filler Application

This application provides a GUI interface for filling Word document templates
with user-provided values. It supports placeholders in the format {{placeholder}}
and handles both text in paragraphs and tables.

Dependencies:
    - tkinter: For the GUI interface
    - python-docx: For handling Word documents
    - re: For regular expression operations
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
import re

class ScrollableFrame(ttk.Frame):
    """
    A custom frame that implements scrolling functionality.
    
    This class creates a scrollable frame that can contain multiple widgets
    and automatically shows a scrollbar when the content exceeds the frame size.
    """
    
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        canvas = tk.Canvas(self, width=380)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

class WordTemplateApp:
    """
    Main application class for the Word Template Filler.
    
    This class handles the main application logic including:
    - Template file selection
    - Placeholder identification
    - User input collection
    - Template filling
    - Saving filled documents
    
    Attributes:
        root (tk.Tk): The main window of the application
        template_path (str): Path to the selected template file
        entries (dict): Dictionary mapping placeholders to their entry widgets
        placeholders (list): List of identified placeholders in the template
    """

    def __init__(self, root):
        """
        Initialize the application.
        
        Args:
            root (tk.Tk): The main window instance
        """
        self.root = root
        self.template_path = None
        self.entries = {}
        self.placeholders = []  # Will be populated from template
        self.create_ui()

    def create_ui(self):
        """
        Set up the GUI components with improved layout.
        
        Creates and arranges all UI elements including:
        - Title
        - Control buttons
        - Scrollable entry area
        - Status bar
        """
        style = ttk.Style()
        style.configure("Action.TButton", padding=5, font=('Helvetica', 10))
        style.configure("Title.TLabel", font=('Helvetica', 12, 'bold'))

        # Main container
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Title
        title = ttk.Label(main_frame, text="Word Template Filler", style="Title.TLabel")
        title.pack(pady=10)

        # Button frame
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=5)

        # Template selection
        button_select = ttk.Button(button_frame, text="Abrir o Template", 
                                 style="Action.TButton", command=self.select_template)
        button_select.pack(side=tk.LEFT, padx=5)

        # Identify placeholders
        button_identify = ttk.Button(button_frame, text="Extrair os Dados a Preencher", 
                                   style="Action.TButton", command=self.identify_placeholders)
        button_identify.pack(side=tk.LEFT, padx=5)

        # Clear entries
        button_clear = ttk.Button(button_frame, text="Limpar os Campos", 
                                style="Action.TButton", command=self.clear_entries)
        button_clear.pack(side=tk.LEFT, padx=5)

        # Scrollable frame for entries
        self.scroll_frame = ScrollableFrame(main_frame)
        self.scroll_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        self.entries_frame = self.scroll_frame.scrollable_frame

        # Create entry fields for predefined placeholders
        self.create_placeholder_entries()

        # Save button frame
        save_frame = ttk.Frame(main_frame)
        save_frame.pack(fill=tk.X, pady=10)

        # Save button
        button_save = ttk.Button(save_frame, text="Salvar o Template Preenchido", 
                               style="Action.TButton", command=self.save_filled_template)
        button_save.pack(side=tk.BOTTOM, pady=5)

        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(fill=tk.X, side=tk.BOTTOM, pady=5)

    def extract_placeholders_in_order(self, doc):
        """
        Extract placeholders while maintaining order of appearance.
        
        Args:
            doc (Document): Word document object
            
        Returns:
            list: Ordered list of unique placeholders found in the document
        """
        ordered_placeholders = []
        placeholder_pattern = r'{{[^}]+}}'

        # Check paragraphs
        for paragraph in doc.paragraphs:
            matches = re.finditer(placeholder_pattern, paragraph.text)
            for match in matches:
                placeholder = match.group(0)
                if placeholder not in ordered_placeholders:
                    ordered_placeholders.append(placeholder)

        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.finditer(placeholder_pattern, cell.text)
                    for match in matches:
                        placeholder = match.group(0)
                        if placeholder not in ordered_placeholders:
                            ordered_placeholders.append(placeholder)

        return ordered_placeholders

    def identify_placeholders(self):
        """
        Identify placeholders in the selected Word template.
        
        Reads the template file and extracts all placeholders in the format {{placeholder}}.
        Creates entry fields for each unique placeholder found.
        
        Displays error messages if no template is selected or if any errors occur during processing.
        """
        if not self.template_path:
            self.status_var.set("Error: No template selected")
            messagebox.showerror("Error", "Please select a template file first.")
            return

        try:
            doc = Document(self.template_path)
            self.placeholders = self.extract_placeholders_in_order(doc)

            if not self.placeholders:
                self.status_var.set("No placeholders found in template")
                messagebox.showinfo("Info", "No placeholders found in the template.")
                return

            # Clear existing entries
            for widget in self.entries_frame.winfo_children():
                widget.destroy()
            self.entries.clear()

            # Create new entries
            self.create_placeholder_entries()
            self.status_var.set(f"Found {len(self.placeholders)} placeholders")
            messagebox.showinfo("Success", f"Found {len(self.placeholders)} placeholders in template")

        except Exception as e:
            self.status_var.set("Error identifying placeholders")
            messagebox.showerror("Error", f"Could not identify placeholders: {e}")

    def create_placeholder_entries(self):
        """
        Create entry fields with improved layout.
        
        Creates labeled entry fields for each placeholder in the template.
        Entry fields are numbered and arranged vertically in a scrollable frame.
        """
        for i, placeholder in enumerate(self.placeholders, 1):
            frame = ttk.Frame(self.entries_frame)
            frame.pack(fill=tk.X, pady=2)
            
            # Add number prefix to show order
            label = ttk.Label(frame, text=f"{i}. {placeholder}", width=25)
            label.pack(side=tk.LEFT, padx=5)
            
            entry = ttk.Entry(frame)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
            self.entries[placeholder] = entry

    def select_template(self):
        """
        Select a Word template file with status update.
        
        Opens a file dialog for template selection and updates the status bar
        with the selected file path.
        """
        self.template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.template_path:
            self.status_var.set(f"Template selected: {self.template_path}")
            messagebox.showinfo("Success", "Template file selected successfully!")

    def save_filled_template(self):
        """
        Save the filled template to a new Word document.
        
        Replaces all placeholders in the template with user-provided values
        and saves the result to a new file. Handles both text in paragraphs
        and tables.
        
        Displays error messages if no template is selected or if any errors
        occur during processing.
        """
        if not self.template_path:
            self.status_var.set("Error: No template selected")
            messagebox.showerror("Error", "Please select a template file first.")
            return

        try:
            doc = Document(self.template_path)
            data = {placeholder: entry.get() for placeholder, entry in self.entries.items()}

            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    paragraph.text = paragraph.text.replace(key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            cell.text = cell.text.replace(key, value)

            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Save Filled Document"
            )
            if save_path:
                doc.save(save_path)
                messagebox.showinfo("Success", f"Filled document saved as:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not fill and save the template: {e}")

    def clear_entries(self):
        """
        Clear all entry fields with confirmation.
        
        Prompts user for confirmation before clearing all entry fields.
        Updates status bar after clearing.
        """
        if messagebox.askyesno("Confirm", "Are you sure you want to clear all entries?"):
            for entry in self.entries.values():
                entry.delete(0, tk.END)
            self.status_var.set("All entries cleared")

if __name__ == "__main__":
    """
    Application entry point
    
    Creates the main window and starts the application.
    Configures the initial window size and style theme.
    """
    root = tk.Tk()
    root.title("Word Template Filler")
    root.geometry("500x600")
    
    # Configure style
    style = ttk.Style()
    style.theme_use('clam')  # or 'alt', 'default', 'classic'
    
    app = WordTemplateApp(root)
    root.mainloop()