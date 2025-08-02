"""
Word Template Filler App

This application provides a web interface for filling Word document templates
with user-provided values. It supports placeholders in the format {{placeholder}}
and handles both text in paragraphs and tables.

Dependencies:
    - streamlit: For the web interface
    - python-docx: For handling Word documents
    - re: For regular expression operations
"""

import streamlit as st
import re
from io import BytesIO

# Page Configuration
st.set_page_config(page_title="Word Template Filler", page_icon="📝", layout="centered")

# Custom CSS
st.markdown(
    """
    <style>
    .main {padding: 2rem}
    .stButton button {width: 100%; margin-top: 1rem}
    </style>
""",
    unsafe_allow_html=True,
)

try:
    from docx import Document

    docx_imported = True
except ImportError:
    docx_imported = False
    st.error(
        "The 'python-docx' library is not installed. Please install it using 'pip install python-docx'."
    )


def extract_placeholders_in_order(doc):
    """
    Extract placeholders while maintaining order of appearance.

    Args:
        doc (Document): Word document object

    Returns:
        list: Ordered list of unique placeholders found in the document
    """
    ordered_placeholders = []
    placeholder_pattern = r"{{[^}]+}}"

    # Check paragraphs
    for paragraph in doc.paragraphs:
        matches = re.finditer(placeholder_pattern, paragraph.text)
        for match in matches:
            placeholder = match.group(0)
            if placeholder not in ordered_placeholders:
                ordered_placeholders.append(placeholder)

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.finditer(placeholder_pattern, cell.text)
                for match in matches:
                    placeholder = match.group(0)
                    if placeholder not in ordered_placeholders:
                        ordered_placeholders.append(placeholder)

    return ordered_placeholders


def fill_template(doc, data):
    """
    Fill the template with user-provided data.

    Args:
        doc (Document): Word document object
        data (dict): Dictionary mapping placeholders to their values
    """
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    cell.text = cell.text.replace(key, value)


def extract_document_structure(doc, placeholders):
    """
    Extract the complete document structure for visual display.
    
    Args:
        doc (Document): Word document object
        placeholders (list): List of placeholders found in the document
        
    Returns:
        list: List of document elements (paragraphs, tables) with structure info
    """
    document_elements = []
    placeholder_pattern = r"{{[^}]+}}"
    
    # Process all paragraphs (including those without placeholders)
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        has_placeholders = bool(re.search(placeholder_pattern, text))
        
        # Get paragraph formatting info
        para_style = get_paragraph_style(paragraph)
        
        document_elements.append({
            'type': 'paragraph',
            'index': para_idx,
            'text': text,
            'has_placeholders': has_placeholders,
            'placeholders': re.findall(placeholder_pattern, text) if has_placeholders else [],
            'style': para_style,
            'is_empty': len(text.strip()) == 0
        })
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            'type': 'table',
            'index': table_idx,
            'rows': []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = {'cells': []}
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                has_placeholders = bool(re.search(placeholder_pattern, text))
                
                cell_data = {
                    'text': text,
                    'has_placeholders': has_placeholders,
                    'placeholders': re.findall(placeholder_pattern, text) if has_placeholders else [],
                    'row_idx': row_idx,
                    'cell_idx': cell_idx
                }
                row_data['cells'].append(cell_data)
            table_data['rows'].append(row_data)
        
        document_elements.append(table_data)
    
    return document_elements


def get_paragraph_style(paragraph):
    """
    Extract paragraph styling information.
    
    Args:
        paragraph: Word paragraph object
        
    Returns:
        dict: Style information
    """
    style_info = {
        'is_heading': False,
        'is_bold': False,
        'is_centered': False,
        'font_size': 'normal',
        'style_name': paragraph.style.name if paragraph.style else 'Normal'
    }
    
    # Check if it's a heading
    if paragraph.style and 'heading' in paragraph.style.name.lower():
        style_info['is_heading'] = True
        if 'heading 1' in paragraph.style.name.lower():
            style_info['font_size'] = 'large'
        elif 'heading 2' in paragraph.style.name.lower():
            style_info['font_size'] = 'medium'
    
    # Check for bold text (simplified check)
    if paragraph.runs:
        bold_runs = [run for run in paragraph.runs if run.bold]
        if len(bold_runs) > len(paragraph.runs) / 2:  # Most runs are bold
            style_info['is_bold'] = True
    
    return style_info


def get_ui_text(language):
    """
    Get UI text based on selected language.

    Args:
        language (str): Selected language ('en' or 'pt')

    Returns:
        dict: Dictionary containing UI text in selected language
    """
    text = {
        "en": {
            # Main interface
            "tab_filler": "📝 Template Filler",
            "tab_help": "ℹ️ Help",
            "choose_file": "Upload Word Template",
            "upload_help": "Upload a Word document with {{placeholder}} fields",
            "fields_found": "Found {} fields to fill",
            "doc_type_detected": "Document type detected: **{}**",
            "generate_button": "Generate Document 📄",
            "download_button": "Download Filled Template ⬇️",
            "fill_all_fields": "Please fill all fields",
            "no_placeholders": "No placeholders found in the template. Make sure your document contains fields in {{placeholder}} format.",
            "document_generated": "Document generated!",
            "processing_template": "Processing template...",
            "creating_document": "Creating document...",
            
            # Help section
            "help_header": "How to Use",
            "help_steps_title": "Steps:",
            "help_step1": "**Upload Template**: Upload Word document with {{placeholder}} format",
            "help_step2": "**Document Preview**: View your document in a Word-like interface",
            "help_step3": "**Fill Fields**: Enter values directly in the highlighted placeholder fields",
            "help_step4": "**Live Preview**: See how your entries will appear in the final document",
            "help_step5": "**Generate**: Click 'Generate Document'",
            "help_step6": "**Download**: Get your completed document",
            "document_preview": "📄 Document Preview Interface",
            "tips_title": "💡 Tips",
            "supported_types": "🔍 Supported Document Types",
        },
        "pt": {
            # Main interface
            "tab_filler": "📝 Preenchimento de Templates",
            "tab_help": "ℹ️ Ajuda",
            "choose_file": "Carregar Template Word",
            "upload_help": "Carregue um documento Word com campos {{placeholder}}",
            "fields_found": "Encontrados {} campos para preencher",
            "doc_type_detected": "Tipo de documento detectado: **{}**",
            "generate_button": "Gerar Documento 📄",
            "download_button": "Descarregar Template Preenchido ⬇️",
            "fill_all_fields": "Por favor, preencha todos os campos",
            "no_placeholders": "Não foram encontrados campos para preencher no template. Certifique-se de que o documento contém campos no formato {{placeholder}}.",
            "document_generated": "Documento gerado!",
            "processing_template": "A processar template...",
            "creating_document": "A criar documento...",
            
            # Help section
            "help_header": "Como Usar",
            "help_steps_title": "Passos:",
            "help_step1": "**Carregar Template**: Carregue documento Word com formato {{placeholder}}",
            "help_step2": "**Pré-visualização do Documento**: Veja o documento numa interface semelhante ao Word",
            "help_step3": "**Preencher Campos**: Insira valores diretamente nos campos destacados",
            "help_step4": "**Pré-visualização em Tempo Real**: Veja como as suas entradas aparecerão no documento final",
            "help_step5": "**Gerar**: Clique em 'Gerar Documento'",
            "help_step6": "**Descarregar**: Obtenha o seu documento completo",
            "document_preview": "📄 Interface de Pré-visualização do Documento",
            "tips_title": "💡 Dicas",
            "supported_types": "🔍 Tipos de Documentos Suportados",
        },
    }
    return text[language]


def detect_document_type(doc):
    """
    Detect the type of document based on its content.
    
    Args:
        doc (Document): Word document object
        
    Returns:
        str: Document type identifier
    """
    # Get the first few paragraphs to analyze content
    content = ""
    for para in doc.paragraphs[:20]:  # Check first 20 paragraphs
        content += para.text.lower() + " "
    
    # Define document type patterns
    if any(word in content for word in ["mediação", "imobiliária", "propriedad", "imóvel", "compra", "venda"]):
        return "real_estate"
    elif any(word in content for word in ["trabalho", "emprego", "contrato de trabalho", "trabalhador", "empregador"]):
        return "labor"
    elif any(word in content for word in ["arrendamento", "aluguel", "locação", "inquilino", "senhorio"]):
        return "rental"
    elif any(word in content for word in ["serviços", "prestação", "fornecedor", "cliente"]):
        return "services"
    else:
        return "generic"


def render_document_like_interface(document_elements, placeholders, ui_text):
    """
    Render the document in a Word-like interface with inline editing.
    
    Args:
        document_elements (list): Document structure elements
        placeholders (list): List of placeholders found in the document
        ui_text (dict): UI text dictionary
        
    Returns:
        dict: Dictionary with placeholder values
    """
    values = {}
    
    # Add Word-like styling
    st.markdown("""
    <style>
    .document-container {
        background-color: white;
        border: 1px solid #d1d5db;
        border-radius: 8px;
        padding: 40px;
        margin: 20px 0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        font-family: 'Times New Roman', serif;
        line-height: 1.6;
    }
    .document-title {
        font-size: 24px;
        font-weight: bold;
        text-align: center;
        margin-bottom: 30px;
        color: #1f2937;
    }
    .document-paragraph {
        margin-bottom: 16px;
        text-align: justify;
        color: #374151;
    }
    .document-paragraph-container {
        margin-bottom: 16px;
        padding: 8px;
        border-left: 3px solid #3b82f6;
        background-color: #f8fafc;
        border-radius: 4px;
    }
    .document-heading {
        font-weight: bold;
        margin: 24px 0 16px 0;
        color: #1f2937;
    }
    .document-heading.large {
        font-size: 20px;
    }
    .document-heading.medium {
        font-size: 18px;
    }
    .document-heading-container {
        margin: 24px 0 16px 0;
        padding: 12px;
        border-left: 4px solid #f59e0b;
        background-color: #fffbeb;
        border-radius: 4px;
    }
    .inline-text {
        display: inline-block;
        margin: 4px 0;
        color: #374151;
        font-family: 'Times New Roman', serif;
    }
    .inline-label {
        display: inline-block;
        font-weight: bold;
        color: #1f2937;
        margin: 8px 0;
        padding: 4px 8px;
        background-color: #e0e7ff;
        border-radius: 4px;
        font-size: 14px;
    }
    .placeholder-highlight {
        background-color: #fef3c7;
        padding: 2px 4px;
        border-radius: 4px;
        border: 1px dashed #f59e0b;
        font-family: monospace;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Create document container
    st.markdown('<div class="document-container">', unsafe_allow_html=True)
    
    # Process each document element
    for element in document_elements:
        if element['type'] == 'paragraph':
            render_paragraph_element(element, values, ui_text)
        elif element['type'] == 'table':
            render_table_element(element, values, ui_text)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    return values


def render_paragraph_element(element, values, ui_text):
    """
    Render a paragraph element with inline placeholder editing.
    
    Args:
        element (dict): Paragraph element data
        values (dict): Dictionary to store placeholder values
        ui_text (dict): UI text dictionary
    """
    text = element['text']
    style = element['style']
    
    # Skip empty paragraphs
    if element['is_empty']:
        st.markdown('<div style="height: 16px;"></div>', unsafe_allow_html=True)
        return
    
    # Apply paragraph styling
    css_class = "document-paragraph"
    if style['is_heading']:
        css_class = f"document-heading {style['font_size']}"
    
    # Process text with placeholders
    if element['has_placeholders']:
        # Show paragraph structure with inline inputs
        st.markdown(f'<div class="{css_class}-container">', unsafe_allow_html=True)
        render_text_with_inline_inputs(text, element['placeholders'], values, ui_text)
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="{css_class}">{text}</div>', unsafe_allow_html=True)


def render_table_element(element, values, ui_text):
    """
    Render a table element with inline placeholder editing.
    
    Args:
        element (dict): Table element data
        values (dict): Dictionary to store placeholder values
        ui_text (dict): UI text dictionary
    """
    st.markdown("**Table:**")
    
    # Process each row and cell with placeholders
    for row_idx, row in enumerate(element['rows']):
        # Check if this row has any placeholders
        row_has_placeholders = any(cell['has_placeholders'] for cell in row['cells'])
        
        if row_has_placeholders:
            st.markdown(f"**Row {row_idx + 1}:**")
            
            for cell_idx, cell in enumerate(row['cells']):
                if cell['has_placeholders']:
                    st.markdown(f"*Column {cell_idx + 1}:*")
                    render_text_with_inline_inputs(cell['text'], cell['placeholders'], values, ui_text)
                else:
                    # Display non-placeholder cells for context
                    if cell['text'].strip():
                        st.markdown(f"*Column {cell_idx + 1}:* {cell['text']}")
            
            st.markdown("---")
        else:
            # Display header or non-placeholder rows for context
            if row_idx == 0 or any(cell['text'].strip() for cell in row['cells']):
                row_text = " | ".join([cell['text'] for cell in row['cells'] if cell['text'].strip()])
                if row_text:
                    st.markdown(f"*{row_text}*")


def render_text_with_inline_inputs(text, placeholders, values, ui_text):
    """
    Render text with inline input fields for placeholders.
    
    Args:
        text (str): Text containing placeholders
        placeholders (list): List of placeholders in the text
        values (dict): Dictionary to store placeholder values
        ui_text (dict): UI text dictionary
        
    Returns:
        None: Renders directly to Streamlit
    """
    if not placeholders:
        # No placeholders, just display the text
        st.markdown(f'<div class="document-paragraph">{text}</div>', unsafe_allow_html=True)
        return
    
    # Split text into parts and render each part with inline inputs
    current_text = text
    text_parts = []
    
    # Sort placeholders by their position in the text
    placeholder_positions = []
    for placeholder in placeholders:
        pos = current_text.find(placeholder)
        if pos != -1:
            placeholder_positions.append((pos, placeholder))
    
    placeholder_positions.sort(key=lambda x: x[0])
    
    # Split text around placeholders
    last_pos = 0
    for pos, placeholder in placeholder_positions:
        # Add text before placeholder
        if pos > last_pos:
            text_parts.append(('text', current_text[last_pos:pos]))
        
        # Add placeholder
        text_parts.append(('placeholder', placeholder))
        last_pos = pos + len(placeholder)
    
    # Add remaining text
    if last_pos < len(current_text):
        text_parts.append(('text', current_text[last_pos:]))
    
    # Render parts with inline inputs
    cols = st.columns(len([part for part in text_parts if part[0] == 'placeholder']) + 1)
    col_idx = 0
    
    current_line = ""
    
    for part_type, content in text_parts:
        if part_type == 'text':
            current_line += content
        elif part_type == 'placeholder':
            # Display text before placeholder if any
            if current_line.strip():
                st.markdown(f'<div class="inline-text">{current_line}</div>', unsafe_allow_html=True)
                current_line = ""
            
            # Create inline input field
            clean_label = content.replace('{{', '').replace('}}', '').replace('_', ' ').title()
            
            # Create a container for inline display
            col1, col2 = st.columns([1, 2])
            
            with col1:
                st.markdown(f'<div class="inline-label">📝 {clean_label}:</div>', 
                           unsafe_allow_html=True)
            
            with col2:
                values[content] = st.text_input(
                    label="",
                    placeholder=f"Enter {clean_label.lower()}",
                    key=f"inline_{content}_{hash(text)}",
                    help=f"This will replace {content} in the document",
                    label_visibility="collapsed"
                )
    
    # Display any remaining text
    if current_line.strip():
        st.markdown(f'<div class="inline-text">{current_line}</div>', unsafe_allow_html=True)


def get_ui_text(language):
    """
    Get UI text based on selected language.

    Args:
        language (str): Selected language ('en' or 'pt')

    Returns:
        dict: Dictionary containing UI text in selected language
    """
    text = {
        "en": {
            # Main interface
            "tab_filler": "📝 Template Filler",
            "tab_help": "ℹ️ Help",
            "choose_file": "Upload Word Template",
            "upload_help": "Upload a Word document with {{placeholder}} fields",
            "fields_found": "Found {} fields to fill",
            "doc_type_detected": "Document type detected: **{}**",
            "generate_button": "Generate Document 📄",
            "download_button": "Download Filled Template ⬇️",
            "fill_all_fields": "Please fill all fields",
            "no_placeholders": "No placeholders found in the template. Make sure your document contains fields in {{placeholder}} format.",
            "document_generated": "Document generated!",
            "processing_template": "Processing template...",
            "creating_document": "Creating document...",
            
            # Help section
            "help_header": "How to Use",
            "help_steps_title": "Steps:",
            "help_step1": "**Upload Template**: Upload Word document with {{placeholder}} format",
            "help_step2": "**Document Preview**: View your document in a Word-like interface",
            "help_step3": "**Fill Fields**: Enter values directly in the highlighted placeholder fields",
            "help_step4": "**Live Preview**: See how your entries will appear in the final document",
            "help_step5": "**Generate**: Click 'Generate Document'",
            "help_step6": "**Download**: Get your completed document",
            "document_preview": "� Document Preview Interface",
            "tips_title": "💡 Tips",
            "supported_types": "🔍 Supported Document Types",
        },
        "pt": {
            # Main interface
            "tab_filler": "📝 Preenchimento de Templates",
            "tab_help": "ℹ️ Ajuda",
            "choose_file": "Carregar Template Word",
            "upload_help": "Carregue um documento Word com campos {{placeholder}}",
            "fields_found": "Encontrados {} campos para preencher",
            "doc_type_detected": "Tipo de documento detectado: **{}**",
            "generate_button": "Gerar Documento 📄",
            "download_button": "Descarregar Template Preenchido ⬇️",
            "fill_all_fields": "Por favor, preencha todos os campos",
            "no_placeholders": "Não foram encontrados campos para preencher no template. Certifique-se de que o documento contém campos no formato {{placeholder}}.",
            "document_generated": "Documento gerado!",
            "processing_template": "A processar template...",
            "creating_document": "A criar documento...",
            
            # Help section
            "help_header": "Como Usar",
            "help_steps_title": "Passos:",
            "help_step1": "**Carregar Template**: Carregue documento Word com formato {{placeholder}}",
            "help_step2": "**Pré-visualização do Documento**: Veja o documento numa interface semelhante ao Word",
            "help_step3": "**Preencher Campos**: Insira valores diretamente nos campos destacados",
            "help_step4": "**Pré-visualização em Tempo Real**: Veja como as suas entradas aparecerão no documento final",
            "help_step5": "**Gerar**: Clique em 'Gerar Documento'",
            "help_step6": "**Descarregar**: Obtenha o seu documento completo",
            "document_preview": "� Interface de Pré-visualização do Documento",
            "tips_title": "💡 Dicas",
            "supported_types": "🔍 Tipos de Documentos Suportados",
        },
    }
    return text[language]


def main():
    # Main Layout with Tabs - will be updated after language selection
    st.title("Word Template Filler")
    
    # === LANGUAGE SELECTION ===
    col1, col2, col3 = st.columns([2, 1, 2])
    with col2:
        language = st.selectbox(
            "🌍 Language / Idioma",
            options=["en", "pt"],
            format_func=lambda x: "🇺🇸 English" if x == "en" else "🇵🇹 Português",
            index=1  # Default to Portuguese
        )
    
    ui_text = get_ui_text(language)
    st.divider()
    
    # Create tabs with dynamic text
    tab1, tab2 = st.tabs([ui_text["tab_filler"], ui_text["tab_help"]])

    with tab1:
        uploaded_file = st.file_uploader(
            ui_text["choose_file"],
            type="docx",
            help=ui_text["upload_help"],
        )

        if uploaded_file:
            with st.spinner(ui_text["processing_template"]):
                doc = Document(uploaded_file)
                placeholders = extract_placeholders_in_order(doc)

            if placeholders:
                # Detect document type for user feedback
                doc_type = detect_document_type(doc)
                type_labels = {
                    "real_estate": "🏠 Real Estate Contract",
                    "labor": "💼 Labor Contract", 
                    "rental": "🏘️ Rental Agreement",
                    "services": "🔧 Service Contract",
                    "generic": "📄 Generic Document"
                }
                
                st.success(ui_text["fields_found"].format(len(placeholders)))
                st.info(ui_text["doc_type_detected"].format(type_labels.get(doc_type, '📄 Generic Document')))
                st.info("� **Document Preview** - Fill in the highlighted fields directly in the document")

                # Extract document structure for Word-like display
                document_elements = extract_document_structure(doc, placeholders)
                
                # Render the document-like interface
                values = render_document_like_interface(document_elements, placeholders, ui_text)

                if st.button(ui_text["generate_button"], type="primary"):
                    if all(values.values()):
                        with st.spinner(ui_text["creating_document"]):
                            # Create a new document from the uploaded file
                            new_doc = Document(uploaded_file)
                            fill_template(new_doc, values)
                            bio = BytesIO()
                            new_doc.save(bio)

                            st.success(ui_text["document_generated"])
                            st.download_button(
                                ui_text["download_button"],
                                bio.getvalue(),
                                "filled_template.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    else:
                        st.warning(ui_text["fill_all_fields"])
            else:
                st.warning(ui_text["no_placeholders"])

    with tab2:
        st.header(ui_text["help_header"])
        
        # Steps section
        st.markdown(f"""
        {ui_text["help_steps_title"]}
        1. {ui_text["help_step1"]}
        2. {ui_text["help_step2"]}
        3. {ui_text["help_step3"]}
        4. {ui_text["help_step4"]}
        5. {ui_text["help_step5"]}
        6. {ui_text["help_step6"]}
        """)
        
        st.subheader(ui_text["document_preview"])
        
        # Contextual display content
        if language == "en":
            st.markdown("""
            The application now shows each placeholder within its document context:
            
            **📄 Document Context Display:**
            - **Paragraph Context**: Shows the full sentence/paragraph where the placeholder appears
            - **Surrounding Text**: Displays text before and after for better understanding
            - **Table Context**: For placeholders in tables, shows table position, headers, and related cells
            - **Highlighted Placeholders**: Placeholders are highlighted in **`bold code format`** within the text
            
            **� Context Information Includes:**
            - **Position**: Where the field appears in the document (paragraph vs table)
            - **Surrounding Content**: Related text that helps understand what to fill in
            - **Table Details**: Column headers, row/column position, and table size
            - **Document Flow**: Previous and following content for context
            """)
        else:
            st.markdown("""
            A aplicação agora mostra cada placeholder dentro do seu contexto no documento:
            
            **📄 Exibição do Contexto do Documento:**
            - **Contexto do Parágrafo**: Mostra a frase/parágrafo completo onde o placeholder aparece
            - **Texto Envolvente**: Exibe texto antes e depois para melhor compreensão
            - **Contexto da Tabela**: Para placeholders em tabelas, mostra posição, cabeçalhos e células relacionadas
            - **Placeholders Destacados**: Placeholders são destacados em **`formato de código negrito`** dentro do texto
            
            **� A Informação de Contexto Inclui:**
            - **Posição**: Onde o campo aparece no documento (parágrafo vs tabela)
            - **Conteúdo Envolvente**: Texto relacionado que ajuda a entender o que preencher
            - **Detalhes da Tabela**: Cabeçalhos de coluna, posição linha/coluna e tamanho da tabela
            - **Fluxo do Documento**: Conteúdo anterior e seguinte para contexto
            """)
        
        st.subheader(ui_text["tips_title"])
        
        # Tips section
        if language == "en":
            st.markdown("""
            - **Inline Editing**: Input fields appear directly next to placeholders in the document flow
            - **Visual Structure**: The interface preserves the exact layout and structure of your Word document
            - **Highlighted Sections**: Paragraphs with placeholders are highlighted with colored borders
            - **Clear Labels**: Each field shows as "📝 Field Name:" exactly where it belongs
            - **Table Support**: Table cells with placeholders are clearly identified and organized
            - **Natural Flow**: Fill fields in the same order you would read the document
            - **Required Fields**: All fields must be filled before generating the document
            - **Live Preview**: See how your entries fit into the document structure as you type
            """)
        else:
            st.markdown("""
            - **Edição Inline**: Campos de entrada aparecem diretamente ao lado dos placeholders no fluxo do documento
            - **Estrutura Visual**: A interface preserva o layout e estrutura exatos do seu documento Word
            - **Seções Destacadas**: Parágrafos com placeholders são destacados com bordas coloridas
            - **Rótulos Claros**: Cada campo mostra como "📝 Nome do Campo:" exatamente onde pertence
            - **Suporte a Tabelas**: Células de tabela com placeholders são claramente identificadas e organizadas
            - **Fluxo Natural**: Preencha campos na mesma ordem que leria o documento
            - **Campos Obrigatórios**: Todos os campos devem ser preenchidos antes de gerar o documento
            - **Pré-visualização ao Vivo**: Veja como as suas entradas se encaixam na estrutura do documento enquanto digita
            """)

        st.subheader(ui_text["supported_types"])
        
        # Supported document types
        if language == "en":
            st.markdown("""
            The application recognizes these document types based on content analysis:
            - **Real Estate**: Mediação, imobiliária, propriedade, imóvel, compra, venda
            - **Labor**: Trabalho, emprego, contrato de trabalho, trabalhador, empregador  
            - **Rental**: Arrendamento, aluguel, locação, inquilino, senhorio
            - **Services**: Serviços, prestação, fornecedor, cliente
            - **Generic**: Any document that doesn't match specific patterns
            """)
        else:
            st.markdown("""
            A aplicação reconhece estes tipos de documentos baseados na análise de conteúdo:
            - **Imobiliário**: Mediação, imobiliária, propriedade, imóvel, compra, venda
            - **Trabalho**: Trabalho, emprego, contrato de trabalho, trabalhador, empregador  
            - **Arrendamento**: Arrendamento, aluguel, locação, inquilino, senhorio
            - **Serviços**: Serviços, prestação, fornecedor, cliente
            - **Genérico**: Qualquer documento que não corresponda a padrões específicos
            """)


if __name__ == "__main__":
    main()
