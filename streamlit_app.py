"""
Word Template Filler App

This application provides a web interface for filling Word document templates
with user-provided values. It supports placeholders in the format {{placeholder}}
and handles both text in paragraphs and tables.

Dependencies:
    - streamlit: For the web interface
    - python-docx: For handling Word documents
    - re: For regular expression operations
"""

import streamlit as st
import re
from io import BytesIO

try:
    from docx import Document
    docx_imported = True
except ImportError:
    docx_imported = False
    st.error("The 'python-docx' library is not installed. Please install it using 'pip install python-docx'.")

def extract_placeholders_in_order(doc):
    """
    Extract placeholders while maintaining order of appearance.
    
    Args:
        doc (Document): Word document object
        
    Returns:
        list: Ordered list of unique placeholders found in the document
    """
    ordered_placeholders = []
    placeholder_pattern = r'{{[^}]+}}'

    # Check paragraphs
    for paragraph in doc.paragraphs:
        matches = re.finditer(placeholder_pattern, paragraph.text)
        for match in matches:
            placeholder = match.group(0)
            if placeholder not in ordered_placeholders:
                ordered_placeholders.append(placeholder)

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.finditer(placeholder_pattern, cell.text)
                for match in matches:
                    placeholder = match.group(0)
                    if placeholder not in ordered_placeholders:
                        ordered_placeholders.append(placeholder)

    return ordered_placeholders

def fill_template(doc, data):
    """
    Fill the template with user-provided data.
    
    Args:
        doc (Document): Word document object
        data (dict): Dictionary mapping placeholders to their values
    """
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    cell.text = cell.text.replace(key, value)

def main():
    if not docx_imported:
        return

    st.title("Word Template Filler")

    uploaded_file = st.file_uploader("Choose a Word template file", type="docx")
    if uploaded_file:
        doc = Document(uploaded_file)
        placeholders = extract_placeholders_in_order(doc)

        if placeholders:
            st.write(f"Found {len(placeholders)} placeholders:")
            data = {}
            for placeholder in placeholders:
                data[placeholder] = st.text_input(f"Value for {placeholder}")

            if st.button("Fill Template"):
                fill_template(doc, data)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Download Filled Template",
                    data=buffer,
                    file_name="filled_template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.write("No placeholders found in the template.")

if __name__ == "__main__":
    main()