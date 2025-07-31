"""
Word Template Filler App

This application provides a web interface for filling Word document templates
with user-provided values. It supports placeholders in the format {{placeholder}}
and handles both text in paragraphs and tables.

Dependencies:
    - streamlit: For the web interface
    - python-docx: For handling Word documents
    - re: For regular expression operations
"""

import streamlit as st
import re
from io import BytesIO

# Page Configuration
st.set_page_config(page_title="Word Template Filler", page_icon="📝", layout="centered")

# Custom CSS
st.markdown(
    """
    <style>
    .main {padding: 2rem}
    .stButton button {width: 100%; margin-top: 1rem}
    </style>
""",
    unsafe_allow_html=True,
)

try:
    from docx import Document

    docx_imported = True
except ImportError:
    docx_imported = False
    st.error(
        "The 'python-docx' library is not installed. Please install it using 'pip install python-docx'."
    )


def extract_placeholders_in_order(doc):
    """
    Extract placeholders while maintaining order of appearance.

    Args:
        doc (Document): Word document object

    Returns:
        list: Ordered list of unique placeholders found in the document
    """
    ordered_placeholders = []
    placeholder_pattern = r"{{[^}]+}}"

    # Check paragraphs
    for paragraph in doc.paragraphs:
        matches = re.finditer(placeholder_pattern, paragraph.text)
        for match in matches:
            placeholder = match.group(0)
            if placeholder not in ordered_placeholders:
                ordered_placeholders.append(placeholder)

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.finditer(placeholder_pattern, cell.text)
                for match in matches:
                    placeholder = match.group(0)
                    if placeholder not in ordered_placeholders:
                        ordered_placeholders.append(placeholder)

    return ordered_placeholders


def fill_template(doc, data):
    """
    Fill the template with user-provided data.

    Args:
        doc (Document): Word document object
        data (dict): Dictionary mapping placeholders to their values
    """
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    cell.text = cell.text.replace(key, value)


def extract_document_sections(doc):
    """
    Extract section headers from the document.
    
    Args:
        doc (Document): Word document object
        
    Returns:
        list: List of section headers found in the document
    """
    sections = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # Look for Cláusula patterns with descriptions in parentheses
            if 'cláusula' in text.lower() and '(' in text and ')' in text:
                match = re.search(r'\(([^)]+)\)', text)
                if match:
                    section_title = match.group(1).strip()
                    if len(section_title) > 3:  # Valid section title
                        sections.append(section_title)
            
            # Look for other section patterns (Artigo, Capítulo, etc.)
            elif any(pattern in text.lower() for pattern in ['artigo', 'capítulo', 'secção', 'título']):
                if re.match(r'(artigo|capítulo|secção|título)\s+\d+', text.lower()):
                    # Extract description after dash or colon
                    if '–' in text or '-' in text or ':' in text:
                        parts = re.split(r'[–\-:]', text, 1)
                        if len(parts) > 1:
                            section_title = parts[1].strip()
                            if len(section_title) > 3:
                                sections.append(section_title)
            
            # Look for numbered sections (1., 2., etc.)
            elif re.match(r'\d+\.\s+[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ]', text):
                section_title = re.sub(r'^\d+\.\s+', '', text).strip()
                if len(section_title) > 3 and len(section_title) < 100:
                    sections.append(section_title)
    
    return sections


def detect_document_type(doc):
    """
    Detect the type of document based on its content.
    
    Args:
        doc (Document): Word document object
        
    Returns:
        str: Document type identifier
    """
    # Get the first few paragraphs to analyze content
    content = ""
    for para in doc.paragraphs[:20]:  # Check first 20 paragraphs
        content += para.text.lower() + " "
    
    # Define document type patterns
    if any(word in content for word in ["mediação", "imobiliária", "propriedad", "imóvel", "compra", "venda"]):
        return "real_estate"
    elif any(word in content for word in ["trabalho", "emprego", "contrato de trabalho", "trabalhador", "empregador"]):
        return "labor"
    elif any(word in content for word in ["arrendamento", "aluguel", "locação", "inquilino", "senhorio"]):
        return "rental"
    elif any(word in content for word in ["serviços", "prestação", "fornecedor", "cliente"]):
        return "services"
    else:
        return "generic"


def create_dynamic_sections(placeholders, doc):
    """
    Create dynamic sections based on document content and placeholders.
    
    Args:
        placeholders (list): List of placeholders found in the document
        doc (Document): Word document object
        
    Returns:
        dict: Dictionary with section names as keys and lists of placeholders as values
    """
    # Extract actual document sections
    document_sections = extract_document_sections(doc)
    
    # Get document type for fallback patterns
    doc_type = detect_document_type(doc)
    fallback_patterns = get_section_templates(doc_type)
    
    # If we found document sections, try to map placeholders to them
    if document_sections:
        sections = {}
        used_placeholders = set()
        
        # For each document section, find related placeholders
        for section_title in document_sections:
            section_placeholders = []
            section_lower = section_title.lower()
            
            # Find placeholders that might belong to this section based on keywords
            for placeholder in placeholders:
                if placeholder in used_placeholders:
                    continue
                    
                placeholder_lower = placeholder.lower()
                
                # Check if placeholder relates to section content
                section_keywords = extract_keywords_from_section(section_lower)
                if any(keyword in placeholder_lower for keyword in section_keywords):
                    section_placeholders.append(placeholder)
                    used_placeholders.add(placeholder)
            
            if section_placeholders:
                sections[section_title] = section_placeholders
        
        # Add remaining placeholders to fallback sections
        remaining_placeholders = [p for p in placeholders if p not in used_placeholders]
        if remaining_placeholders:
            # Use fallback categorization for remaining placeholders
            fallback_sections = categorize_with_fallback(remaining_placeholders, fallback_patterns)
            sections.update(fallback_sections)
        
        return sections
    
    # If no document sections found, use fallback categorization
    return categorize_with_fallback(placeholders, fallback_patterns)


def extract_keywords_from_section(section_text):
    """
    Extract relevant keywords from section text to match with placeholders.
    
    Args:
        section_text (str): Section title text
        
    Returns:
        list: List of keywords for matching
    """
    # Enhanced mapping based on actual document sections
    keyword_mappings = {
        'identificação': ['nome', 'id', 'fiscal', 'cliente', 'estado_civil', 'regime', 'morada_cliente', 'codpostal', 'telemóvel', 'email'],
        'imóvel': ['fracçao', 'morada_do_imov', 'lugar', 'area', 'terreno', 'bruta', 'registo', 'conservatoria', 'matricial', 'nip', 'licen', 'camara'],
        'negócio': ['valor', 'preço', 'extenso'],
        'remuneração': ['percentagem', 'valor'],
        'contratação': ['regime'],
        'angariador': ['angariador', 'cartao', 'cidadao', 'nif'],
        'prazo': ['dia', 'mes', 'ano', 'data'],
        'duração': ['dia', 'mes', 'ano', 'data'],
        'colaboração': ['cliente', 'contratante'],
        'dados': ['nome', 'contacto', 'email'],
        'ónus': ['valor', 'encargo'],
        'garantias': ['atividade', 'mediação']
    }
    
    keywords = []
    section_lower = section_text.lower()
    
    for concept, related_keywords in keyword_mappings.items():
        if concept in section_lower:
            keywords.extend(related_keywords)
    
    # Add partial word matches
    for word in section_text.lower().split():
        if len(word) > 3:  # Only meaningful words
            keywords.append(word)
    
    return keywords


def categorize_with_fallback(placeholders, fallback_patterns):
    """
    Categorize placeholders using fallback patterns.
    
    Args:
        placeholders (list): List of placeholders to categorize
        fallback_patterns (dict): Fallback section patterns
        
    Returns:
        dict: Categorized sections
    """
    sections = {section_name: [] for section_name in fallback_patterns.keys()}
    if "Other" not in sections:
        sections["Other"] = []
    
    for placeholder in placeholders:
        placeholder_lower = placeholder.lower()
        categorized = False
        
        # Check each section pattern
        for section_name, patterns in fallback_patterns.items():
            for pattern in patterns:
                if pattern in placeholder_lower:
                    sections[section_name].append(placeholder)
                    categorized = True
                    break
            if categorized:
                break
        
        # If not categorized, put in "Other"
        if not categorized:
            sections["Other"].append(placeholder)
    
    # Remove empty sections
    return {k: v for k, v in sections.items() if v}


def get_section_templates(document_type):
    """
    Get section templates based on document type.
    
    Args:
        document_type (str): Type of document detected
        
    Returns:
        dict: Section patterns for the document type
    """
    templates = {
        "real_estate": {
            "Client Information": ["cliente", "nome", "estado_civil", "regime", "morada_cliente", "codpostal-cidade", 
                                 "id", "fiscal", "telemóvel", "email"],
            "Property Details": ["fracçao", "morada_do_imov", "lugar", "codpostal_imovel", "area", "terreno", "bruta"],
            "Legal & Registry": ["registo", "conservatoria", "artigo", "matricial", "nip", "licen", "camara", "municipal", "data"],
            "Financial Terms": ["valor", "percentagem", "regime", "mediacao"],
            "Real Estate Agent": ["angariador", "cartao", "cidadao", "nif"],
            "Contract Date": ["dia", "mes", "ano"]
        },
        "labor": {
            "Employee Information": ["nome", "trabalhador", "empregado", "morada", "id", "fiscal", "nascimento"],
            "Employer Information": ["empresa", "empregador", "sede", "nipc", "atividade"],
            "Job Details": ["cargo", "função", "categoria", "local", "horário", "trabalho"],
            "Compensation": ["salário", "remuneração", "vencimento", "valor", "subsídio"],
            "Contract Terms": ["duração", "período", "prazo", "início", "fim"],
            "Date & Signatures": ["dia", "mes", "ano", "assinatura"]
        },
        "rental": {
            "Tenant Information": ["inquilino", "arrendatário", "nome", "morada", "id", "fiscal"],
            "Landlord Information": ["senhorio", "proprietário", "arrendador"],
            "Property Information": ["imóvel", "fracção", "morada", "area", "tipologia"],
            "Rental Terms": ["renda", "valor", "duração", "prazo", "caução"],
            "Contract Date": ["dia", "mes", "ano"]
        },
        "services": {
            "Client Information": ["cliente", "nome", "empresa", "morada", "contacto"],
            "Service Provider": ["prestador", "fornecedor", "empresa"],
            "Service Details": ["serviço", "descrição", "objeto", "atividade"],
            "Financial Terms": ["valor", "preço", "pagamento", "faturação"],
            "Contract Date": ["dia", "mes", "ano"]
        },
        "generic": {
            "Personal Information": ["nome", "morada", "contacto", "email", "telefone"],
            "Financial Information": ["valor", "preço", "custo", "pagamento"],
            "Dates": ["data", "dia", "mes", "ano"],
            "Other Information": []  # Will catch remaining fields
        }
    }
    
    return templates.get(document_type, templates["generic"])


def get_ui_text(language):
    """
    Get UI text based on selected language.

    Args:
        language (str): Selected language ('en' or 'pt')

    Returns:
        dict: Dictionary containing UI text in selected language
    """
    text = {
        "en": {
            # Main interface
            "tab_filler": "📝 Template Filler",
            "tab_help": "ℹ️ Help",
            "choose_file": "Upload Word Template",
            "upload_help": "Upload a Word document with {{placeholder}} fields",
            "fields_found": "Found {} fields to fill",
            "doc_type_detected": "Document type detected: **{}**",
            "using_doc_sections": "📋 Using **{} document-specific sections** from template structure",
            "using_smart_categorization": "📋 Using **smart categorization** based on field content",
            "generate_button": "Generate Document 📄",
            "download_button": "Download Filled Template ⬇️",
            "fill_all_fields": "Please fill all fields",
            "no_placeholders": "No placeholders found in the template. Make sure your document contains fields in {{placeholder}} format.",
            "document_generated": "Document generated!",
            "processing_template": "Processing template...",
            "creating_document": "Creating document...",
            
            # Help section
            "help_header": "How to Use",
            "help_steps_title": "Steps:",
            "help_step1": "**Upload Template**: Upload Word document with {{placeholder}} format",
            "help_step2": "**Fill Fields**: Enter values for each placeholder organized by sections",
            "help_step3": "**Generate**: Click 'Generate Document'",
            "help_step4": "**Download**: Get your completed document",
            "smart_organization": "📋 Smart Field Organization",
            "tips_title": "💡 Tips",
            "supported_types": "🔍 Supported Document Types",
        },
        "pt": {
            # Main interface
            "tab_filler": "📝 Preenchimento de Templates",
            "tab_help": "ℹ️ Ajuda",
            "choose_file": "Carregar Template Word",
            "upload_help": "Carregue um documento Word com campos {{placeholder}}",
            "fields_found": "Encontrados {} campos para preencher",
            "doc_type_detected": "Tipo de documento detectado: **{}**",
            "using_doc_sections": "📋 Usando **{} seções específicas do documento** da estrutura do template",
            "using_smart_categorization": "📋 Usando **categorização inteligente** baseada no conteúdo dos campos",
            "generate_button": "Gerar Documento 📄",
            "download_button": "Descarregar Template Preenchido ⬇️",
            "fill_all_fields": "Por favor, preencha todos os campos",
            "no_placeholders": "Não foram encontrados campos para preencher no template. Certifique-se de que o documento contém campos no formato {{placeholder}}.",
            "document_generated": "Documento gerado!",
            "processing_template": "A processar template...",
            "creating_document": "A criar documento...",
            
            # Help section
            "help_header": "Como Usar",
            "help_steps_title": "Passos:",
            "help_step1": "**Carregar Template**: Carregue documento Word com formato {{placeholder}}",
            "help_step2": "**Preencher Campos**: Insira valores para cada placeholder organizados por seções",
            "help_step3": "**Gerar**: Clique em 'Gerar Documento'",
            "help_step4": "**Descarregar**: Obtenha o seu documento completo",
            "smart_organization": "📋 Organização Inteligente de Campos",
            "tips_title": "💡 Dicas",
            "supported_types": "🔍 Tipos de Documentos Suportados",
        },
    }
    return text[language]


def main():
    # Main Layout with Tabs - will be updated after language selection
    st.title("Word Template Filler")
    
    # === LANGUAGE SELECTION ===
    col1, col2, col3 = st.columns([2, 1, 2])
    with col2:
        language = st.selectbox(
            "🌍 Language / Idioma",
            options=["en", "pt"],
            format_func=lambda x: "🇺🇸 English" if x == "en" else "🇵🇹 Português",
            index=1  # Default to Portuguese
        )
    
    ui_text = get_ui_text(language)
    st.divider()
    
    # Create tabs with dynamic text
    tab1, tab2 = st.tabs([ui_text["tab_filler"], ui_text["tab_help"]])

    with tab1:
        uploaded_file = st.file_uploader(
            ui_text["choose_file"],
            type="docx",
            help=ui_text["upload_help"],
        )

        if uploaded_file:
            with st.spinner(ui_text["processing_template"]):
                doc = Document(uploaded_file)
                placeholders = extract_placeholders_in_order(doc)

            if placeholders:
                # Detect document type for user feedback
                doc_type = detect_document_type(doc)
                type_labels = {
                    "real_estate": "🏠 Real Estate Contract",
                    "labor": "💼 Labor Contract", 
                    "rental": "🏘️ Rental Agreement",
                    "services": "🔧 Service Contract",
                    "generic": "📄 Generic Document"
                }
                
                # Check if we found document-specific sections
                document_sections = extract_document_sections(doc)
                
                st.success(ui_text["fields_found"].format(len(placeholders)))
                st.info(ui_text["doc_type_detected"].format(type_labels.get(doc_type, '📄 Generic Document')))
                
                if document_sections:
                    st.info(ui_text["using_doc_sections"].format(len(document_sections)))
                else:
                    st.info(ui_text["using_smart_categorization"])

                # Categorize placeholders into sections based on document content
                sections = create_dynamic_sections(placeholders, doc)
                values = {}

                # Display fields organized by sections
                for section_name, section_placeholders in sections.items():
                    if section_placeholders:  # Only show sections with placeholders
                        st.subheader(f"📋 {section_name}")
                        
                        # Create columns for this section
                        col1, col2 = st.columns(2)
                        
                        for i, placeholder in enumerate(section_placeholders):
                            with col1 if i % 2 == 0 else col2:
                                # Create a cleaner label by removing {{ and }}
                                clean_label = placeholder.replace('{{', '').replace('}}', '').replace('_', ' ').title()
                                values[placeholder] = st.text_input(
                                    clean_label,
                                    placeholder=f"Enter {clean_label.lower()}",
                                    key=f"field_{placeholder}"
                                )
                        
                        st.divider()  # Add visual separation between sections

                if st.button(ui_text["generate_button"], type="primary"):
                    if all(values.values()):
                        with st.spinner(ui_text["creating_document"]):
                            # Create a new document from the uploaded file
                            new_doc = Document(uploaded_file)
                            fill_template(new_doc, values)
                            bio = BytesIO()
                            new_doc.save(bio)

                            st.success(ui_text["document_generated"])
                            st.download_button(
                                ui_text["download_button"],
                                bio.getvalue(),
                                "filled_template.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    else:
                        st.warning(ui_text["fill_all_fields"])
            else:
                st.warning(ui_text["no_placeholders"])

    with tab2:
        st.header(ui_text["help_header"])
        
        # Steps section
        st.markdown(f"""
        {ui_text["help_steps_title"]}
        1. {ui_text["help_step1"]}
        2. {ui_text["help_step2"]}
        3. {ui_text["help_step3"]}
        4. {ui_text["help_step4"]}
        """)
        
        st.subheader(ui_text["smart_organization"])
        
        # Smart organization content (keeping this in both languages for now)
        if language == "en":
            st.markdown("""
            The application automatically detects the document type and organizes fields accordingly:
            
            **🏠 Real Estate Contracts:**
            - Client Information, Property Details, Legal & Registry, Financial Terms, Real Estate Agent, Contract Date
            
            **💼 Labor Contracts:**
            - Employee Information, Employer Information, Job Details, Compensation, Contract Terms, Date & Signatures
            
            **🏘️ Rental Agreements:**
            - Tenant Information, Landlord Information, Property Information, Rental Terms, Contract Date
            
            **🔧 Service Contracts:**
            - Client Information, Service Provider, Service Details, Financial Terms, Contract Date
            
            **📄 Generic Documents:**
            - Personal Information, Financial Information, Dates, Other Information
            """)
        else:
            st.markdown("""
            A aplicação detecta automaticamente o tipo de documento e organiza os campos em conformidade:
            
            **🏠 Contratos Imobiliários:**
            - Informação do Cliente, Detalhes da Propriedade, Registo Legal, Termos Financeiros, Agente Imobiliário, Data do Contrato
            
            **💼 Contratos de Trabalho:**
            - Informação do Funcionário, Informação do Empregador, Detalhes do Trabalho, Compensação, Termos do Contrato, Data e Assinaturas
            
            **🏘️ Contratos de Arrendamento:**
            - Informação do Inquilino, Informação do Senhorio, Informação da Propriedade, Termos de Arrendamento, Data do Contrato
            
            **🔧 Contratos de Serviços:**
            - Informação do Cliente, Prestador de Serviços, Detalhes do Serviço, Termos Financeiros, Data do Contrato
            
            **📄 Documentos Genéricos:**
            - Informação Pessoal, Informação Financeira, Datas, Outras Informações
            """)
        
        st.subheader(ui_text["tips_title"])
        
        # Tips section
        if language == "en":
            st.markdown("""
            - **Automatic Detection**: Document type is detected from content keywords
            - **Clean Labels**: Field labels are automatically cleaned ({{field_name}} becomes "Field Name")
            - **Required Fields**: All fields must be filled before generating the document
            - **Format Preservation**: Original template structure and formatting is preserved
            - **Multi-Language**: Supports Portuguese, English, and other languages
            - **File Format**: Download completed documents in .docx format
            """)
        else:
            st.markdown("""
            - **Detecção Automática**: O tipo de documento é detectado através de palavras-chave do conteúdo
            - **Rótulos Limpos**: Os rótulos dos campos são automaticamente limpos ({{nome_campo}} torna-se "Nome Campo")
            - **Campos Obrigatórios**: Todos os campos devem ser preenchidos antes de gerar o documento
            - **Preservação do Formato**: A estrutura e formatação original do template é preservada
            - **Multi-Idioma**: Suporta Português, Inglês e outros idiomas
            - **Formato de Ficheiro**: Descarregue documentos completos em formato .docx
            """)

        st.subheader(ui_text["supported_types"])
        
        # Supported document types
        if language == "en":
            st.markdown("""
            The application recognizes these document types based on content analysis:
            - **Real Estate**: Mediação, imobiliária, propriedade, imóvel, compra, venda
            - **Labor**: Trabalho, emprego, contrato de trabalho, trabalhador, empregador  
            - **Rental**: Arrendamento, aluguel, locação, inquilino, senhorio
            - **Services**: Serviços, prestação, fornecedor, cliente
            - **Generic**: Any document that doesn't match specific patterns
            """)
        else:
            st.markdown("""
            A aplicação reconhece estes tipos de documentos baseados na análise de conteúdo:
            - **Imobiliário**: Mediação, imobiliária, propriedade, imóvel, compra, venda
            - **Trabalho**: Trabalho, emprego, contrato de trabalho, trabalhador, empregador  
            - **Arrendamento**: Arrendamento, aluguel, locação, inquilino, senhorio
            - **Serviços**: Serviços, prestação, fornecedor, cliente
            - **Genérico**: Qualquer documento que não corresponda a padrões específicos
            """)


if __name__ == "__main__":
    main()
