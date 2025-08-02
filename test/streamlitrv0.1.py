"""
Word Template Filler App

This application provides a web interface for filling Word document templates
with user-provided values. It supports placeholders in the format {{placeholder}}
and handles both text in paragraphs and tables.

Dependencies:
    - streamlit: For the web interface
    - python-docx: For handling Word documents
    - re: For regular expression operations
"""

import streamlit as st
import re
from io import BytesIO

try:
    from docx import Document
    docx_imported = True
except ImportError:
    docx_imported = False
    st.error("The 'python-docx' library is not installed. Please install it using 'pip install python-docx'.")

def extract_placeholders_in_order(doc):
    """
    Extract placeholders while maintaining order of appearance.
    
    Args:
        doc (Document): Word document object
        
    Returns:
        list: Ordered list of unique placeholders found in the document
    """
    ordered_placeholders = []
    placeholder_pattern = r'{{[^}]+}}'

    # Check paragraphs
    for paragraph in doc.paragraphs:
        matches = re.finditer(placeholder_pattern, paragraph.text)
        for match in matches:
            placeholder = match.group(0)
            if placeholder not in ordered_placeholders:
                ordered_placeholders.append(placeholder)

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.finditer(placeholder_pattern, cell.text)
                for match in matches:
                    placeholder = match.group(0)
                    if placeholder not in ordered_placeholders:
                        ordered_placeholders.append(placeholder)

    return ordered_placeholders

def fill_template(doc, data):
    """
    Fill the template with user-provided data.
    
    Args:
        doc (Document): Word document object
        data (dict): Dictionary mapping placeholders to their values
    """
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    cell.text = cell.text.replace(key, value)

def get_ui_text(language):
    """
    Get UI text based on selected language.
    
    Args:
        language (str): Selected language ('en' or 'pt')
        
    Returns:
        dict: Dictionary containing UI text in selected language
    """
    text = {
        'en': {
            'title': "Word Template Filling Form",
            'description': "Choose a template file with {{placeholder}} fields, the application will extract the fields to fill. Fill in the fields and save the completed file.",
            'choose_file': "Choose Word Document",
            'fields_found': "Found {} fields to fill:",
            'field_to_fill': "Field to fill",
            'create_file': "Create File",
            'download_file': "Download File",
            'no_placeholders': "No placeholders found in the template."
        },
        'pt': {
            'title': "Formulário de Preenchimento - Templates Word",
            'description': "Escolher o ficheiro template onde campos {{placeholder}} estao inseridos, a aplicação extrai os campos a preencher. Preencha os campos e guarde o ficheiro preenchido.",
            'choose_file': "Escolher o Documento Word",
            'fields_found': "Encontrados {} espaços a preencher:",
            'field_to_fill': "Campo a preencher",
            'create_file': "Criar Ficheiro",
            'download_file': "Download de Ficheiro",
            'no_placeholders': "Não foram encontrados campos para preencher no template."
        }
    }
    return text[language]

def main():
    if not docx_imported:
        return

    # Language selection
    language = st.selectbox(
        "Select Language / Selecione o Idioma",
        options=['en', 'pt'],
        format_func=lambda x: "English" if x == 'en' else "Português"
    )
    
    ui_text = get_ui_text(language)
    
    st.title(ui_text['title'])
    st.markdown(ui_text['description'])

    uploaded_file = st.file_uploader(ui_text['choose_file'], type="docx")
    if uploaded_file:
        doc = Document(uploaded_file)
        placeholders = extract_placeholders_in_order(doc)

        if placeholders:
            st.write(ui_text['fields_found'].format(len(placeholders)))
            data = {}
            for placeholder in placeholders:
                data[placeholder] = st.text_input(f"{ui_text['field_to_fill']} {placeholder}")

            if st.button(ui_text['create_file']):
                fill_template(doc, data)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label=ui_text['download_file'],
                    data=buffer,
                    file_name="filled_template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.write(ui_text['no_placeholders'])

if __name__ == "__main__":
    main()